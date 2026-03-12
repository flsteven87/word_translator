from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.models.translation import (
    ParagraphStyle,
    TranslationDirection,
    TranslationResult,
)

_NON_EXPORTABLE_STYLES = frozenset({ParagraphStyle.FIGURE, ParagraphStyle.TABLE})

_HEADING_STYLES = frozenset({
    ParagraphStyle.TITLE,
    ParagraphStyle.HEADING_1,
    ParagraphStyle.HEADING_2,
    ParagraphStyle.HEADING_3,
    ParagraphStyle.HEADING_4,
})


class WordExporter:
    def export(self, result: TranslationResult) -> bytes:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)

        if result.direction == TranslationDirection.ZH_TO_EN:
            left_header, right_header = "中文（原文）", "English (Translation)"
        else:
            left_header, right_header = "English (Original)", "中文（翻譯）"

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        header = table.rows[0].cells
        header[0].text = left_header
        header[1].text = right_header
        for cell in header:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        for para in result.paragraphs:
            if para.style in _NON_EXPORTABLE_STYLES:
                continue
            row = table.add_row().cells
            is_heading = para.style in _HEADING_STYLES
            row[0].text = para.original
            row[1].text = para.translated
            if is_heading:
                for cell in row:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        # Force equal column widths by disabling autofit and setting fixed layout
        table.autofit = False
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = tbl_pr.makeelement(qn("w:tblLayout"), {})
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

        col_width = Inches(3.25)
        for col in table.columns:
            col.width = col_width

        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(4)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
