from io import BytesIO
from unittest.mock import MagicMock, patch

from docx import Document

from src.models.translation import ParagraphStyle
from src.services.document_parser import DocumentParser, ParsedParagraph, _parse_markdown


def _make_docx(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_chunk(chunk_type: str, markdown: str) -> MagicMock:
    chunk = MagicMock()
    chunk.type = chunk_type
    chunk.markdown = markdown
    return chunk


def test_parse_extracts_paragraphs():
    content = _make_docx(["First paragraph.", "Second paragraph."])
    result = DocumentParser(vision_agent_api_key="test-key").parse(content, "test.docx")
    assert len(result) == 2
    assert result[0] == ParsedParagraph(text="First paragraph.", style=ParagraphStyle.NORMAL)
    assert result[1] == ParsedParagraph(text="Second paragraph.", style=ParagraphStyle.NORMAL)


def test_parse_skips_empty_paragraphs():
    content = _make_docx(["Hello.", "", "  ", "World."])
    result = DocumentParser(vision_agent_api_key="test-key").parse(content, "test.docx")
    assert len(result) == 2
    assert result[0].text == "Hello."
    assert result[1].text == "World."


def test_parse_extracts_headings():
    doc = Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Body text.")
    buf = BytesIO()
    doc.save(buf)
    result = DocumentParser(vision_agent_api_key="test-key").parse(buf.getvalue(), "test.docx")
    assert len(result) == 2
    assert result[0] == ParsedParagraph(text="Title", style=ParagraphStyle.HEADING_1)
    assert result[1] == ParsedParagraph(text="Body text.", style=ParagraphStyle.NORMAL)


def test_parse_pdf_uses_ade_when_available():
    fake_response = MagicMock()
    fake_response.chunks = [
        _make_chunk("chunkTitle", "# Title"),
        _make_chunk("chunkText", "First paragraph.\n\nSecond paragraph."),
    ]

    parser = DocumentParser(vision_agent_api_key="test-key")
    with patch.object(parser._ade_client, "parse", return_value=fake_response) as mock_parse:
        result = parser.parse(b"fake-pdf-bytes", "test.pdf")

    mock_parse.assert_called_once()
    assert len(result) == 3
    assert result[0] == ParsedParagraph(text="Title", style=ParagraphStyle.TITLE)
    assert result[1] == ParsedParagraph(text="First paragraph.", style=ParagraphStyle.NORMAL)
    assert result[2] == ParsedParagraph(text="Second paragraph.", style=ParagraphStyle.NORMAL)


def test_parse_pdf_falls_back_to_pymupdf_on_ade_failure():
    parser = DocumentParser(vision_agent_api_key="test-key")

    with patch.object(parser._ade_client, "parse", side_effect=Exception("API down")), \
         patch("src.services.document_parser.pymupdf4llm") as mock_pymupdf4llm, \
         patch("src.services.document_parser.pymupdf") as mock_pymupdf:
        mock_doc = MagicMock()
        mock_doc.page_count = 1
        mock_doc.__enter__ = MagicMock(return_value=mock_doc)
        mock_doc.__exit__ = MagicMock(return_value=False)
        mock_pymupdf.open.return_value = mock_doc
        mock_pymupdf4llm.IdentifyHeaders.return_value = {}
        mock_pymupdf4llm.to_markdown.return_value = "# Fallback Title\n\nFallback body."

        result = parser.parse(b"fake-pdf-bytes", "test.pdf")

    assert len(result) == 2
    assert result[0].text == "Fallback Title"
    assert result[1].text == "Fallback body."


def test_parse_pdf_falls_back_when_ade_returns_empty_chunks():
    fake_response = MagicMock()
    fake_response.chunks = []

    parser = DocumentParser(vision_agent_api_key="test-key")
    with patch.object(parser._ade_client, "parse", return_value=fake_response), \
         patch("src.services.document_parser.pymupdf4llm") as mock_pymupdf4llm, \
         patch("src.services.document_parser.pymupdf") as mock_pymupdf:
        mock_doc = MagicMock()
        mock_doc.page_count = 1
        mock_doc.__enter__ = MagicMock(return_value=mock_doc)
        mock_doc.__exit__ = MagicMock(return_value=False)
        mock_pymupdf.open.return_value = mock_doc
        mock_pymupdf4llm.IdentifyHeaders.return_value = {}
        mock_pymupdf4llm.to_markdown.return_value = "# Fallback"

        result = parser.parse(b"fake-pdf-bytes", "test.pdf")

    assert len(result) == 1
    assert result[0].text == "Fallback"


# --- ADE chunk-based parsing tests ---


def test_chunks_figure_identified():
    fake_response = MagicMock()
    fake_response.chunks = [
        _make_chunk("chunkFigure", "<::bar chart::>Y-axis label"),
    ]
    parser = DocumentParser(vision_agent_api_key="test-key")
    with patch.object(parser._ade_client, "parse", return_value=fake_response):
        result = parser.parse(b"fake-pdf", "test.pdf")

    assert len(result) == 1
    assert result[0].style == ParagraphStyle.FIGURE
    assert "<::bar chart::>" in result[0].text


def test_chunks_table_identified():
    fake_response = MagicMock()
    fake_response.chunks = [
        _make_chunk("chunkTable", '<table id="t1"><tr><td>A</td></tr></table>'),
    ]
    parser = DocumentParser(vision_agent_api_key="test-key")
    with patch.object(parser._ade_client, "parse", return_value=fake_response):
        result = parser.parse(b"fake-pdf", "test.pdf")

    assert len(result) == 1
    assert result[0].style == ParagraphStyle.TABLE


def test_chunks_skip_marginalia():
    fake_response = MagicMock()
    fake_response.chunks = [
        _make_chunk("chunkMarginalia", "Page note"),
        _make_chunk("chunkText", "Real content."),
    ]
    parser = DocumentParser(vision_agent_api_key="test-key")
    with patch.object(parser._ade_client, "parse", return_value=fake_response):
        result = parser.parse(b"fake-pdf", "test.pdf")

    assert len(result) == 1
    assert result[0].text == "Real content."
    assert result[0].style == ParagraphStyle.NORMAL


def test_chunks_preserve_headings_in_text():
    fake_response = MagicMock()
    fake_response.chunks = [
        _make_chunk("chunkTitle", "## Section Heading"),
    ]
    parser = DocumentParser(vision_agent_api_key="test-key")
    with patch.object(parser._ade_client, "parse", return_value=fake_response):
        result = parser.parse(b"fake-pdf", "test.pdf")

    assert len(result) == 1
    assert result[0].style == ParagraphStyle.HEADING_1
    assert result[0].text == "Section Heading"


def test_chunks_logo_maps_to_figure():
    fake_response = MagicMock()
    fake_response.chunks = [
        _make_chunk("chunkLogo", "<::company logo::>"),
    ]
    parser = DocumentParser(vision_agent_api_key="test-key")
    with patch.object(parser._ade_client, "parse", return_value=fake_response):
        result = parser.parse(b"fake-pdf", "test.pdf")

    assert len(result) == 1
    assert result[0].style == ParagraphStyle.FIGURE


def test_chunks_skip_empty_markdown():
    fake_response = MagicMock()
    fake_response.chunks = [
        _make_chunk("chunkText", ""),
        _make_chunk("chunkText", "   "),
        _make_chunk("chunkText", "Actual text."),
    ]
    parser = DocumentParser(vision_agent_api_key="test-key")
    with patch.object(parser._ade_client, "parse", return_value=fake_response):
        result = parser.parse(b"fake-pdf", "test.pdf")

    assert len(result) == 1
    assert result[0].text == "Actual text."


# --- _parse_markdown fallback tests ---


def test_parse_markdown_ade_delimiter_recognized():
    md = "<::bar chart::>Y-axis label: results"
    result = _parse_markdown(md)
    assert len(result) == 1
    assert result[0].style == ParagraphStyle.FIGURE
    assert "<::bar chart::>" in result[0].text


def test_parse_markdown_html_table_single_line():
    md = '<table id="1"><tr><td>Data</td></tr></table>'
    result = _parse_markdown(md)
    assert len(result) == 1
    assert result[0].style == ParagraphStyle.TABLE
    assert "<table" in result[0].text


def test_parse_markdown_html_table_multiline():
    md = (
        "Some intro text.\n"
        "\n"
        '<table id="t1">\n'
        "  <tr><td>Cell A</td></tr>\n"
        "  <tr><td>Cell B</td></tr>\n"
        "</table>\n"
        "\n"
        "After table."
    )
    result = _parse_markdown(md)
    assert len(result) == 3
    assert result[0] == ParsedParagraph(text="Some intro text.", style=ParagraphStyle.NORMAL)
    assert result[1].style == ParagraphStyle.TABLE
    assert "<table" in result[1].text
    assert "Cell A" in result[1].text
    assert "Cell B" in result[1].text
    assert result[2] == ParsedParagraph(text="After table.", style=ParagraphStyle.NORMAL)


def test_parse_markdown_unterminated_html_table_flushed():
    md = '<table id="t1">\n  <tr><td>No close tag</td></tr>'
    result = _parse_markdown(md)
    assert len(result) == 1
    assert result[0].style == ParagraphStyle.TABLE
