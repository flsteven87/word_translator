import asyncio
from io import BytesIO
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document

from src.core.exceptions import NotFoundError
from src.models.translation import ParagraphStyle, TranslationResult
from src.services.translation_service import TranslationService


@pytest.fixture
def service(tmp_path):
    return TranslationService(
        storage_dir=tmp_path,
        openai_api_key="test-key",
        openai_model="gpt-4o-mini",
        vision_agent_api_key="test-key",
    )


def _make_docx(paragraphs: list[str]) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_translate_document(service):
    docx_content = _make_docx(["Hello world.", "Good morning."])

    with patch.object(
        service._strategy, "translate", new_callable=AsyncMock
    ) as mock_translate:
        mock_translate.return_value = ["你好世界。", "早安。"]
        result = await service.translate_document(docx_content, "test.docx")

    assert isinstance(result, TranslationResult)
    assert result.filename == "test.docx"
    assert len(result.paragraphs) == 2
    assert result.paragraphs[0].original == "Hello world."
    assert result.paragraphs[0].translated == "你好世界。"
    assert result.paragraphs[1].original == "Good morning."
    assert result.paragraphs[1].translated == "早安。"


@pytest.mark.asyncio
async def test_translate_document_is_persisted(service):
    docx_content = _make_docx(["Hello."])

    with patch.object(
        service._strategy, "translate", new_callable=AsyncMock
    ) as mock_translate:
        mock_translate.return_value = ["你好。"]
        result = await service.translate_document(docx_content, "test.docx")

    loaded = service.get_translation(str(result.id))
    assert loaded.id == result.id


def test_list_translations_empty(service):
    assert service.list_translations() == []


def test_get_translation_not_found(service):
    with pytest.raises(NotFoundError):
        service.get_translation("nonexistent")


def test_export_translation(service):
    docx_content = _make_docx(["Hello."])

    with patch.object(service._strategy, "translate", new_callable=AsyncMock) as m:
        m.return_value = ["你好。"]
        result = asyncio.get_event_loop().run_until_complete(
            service.translate_document(docx_content, "test.docx")
        )

    loaded = service.get_translation(str(result.id))
    docx_bytes, filename = service.export_translation(loaded)
    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 0
    assert filename == "test_中文.docx"


@pytest.mark.asyncio
async def test_short_paragraphs_each_get_own_translation(service):
    """Regression: short paragraphs must each map 1:1 to their translation.

    Previously, short paragraphs were merged into one chunk and the LLM
    could drop or combine translations, causing misalignment.
    """
    docx_content = _make_docx([
        "The abbreviations (ed.) or (eds.) should follow the Editors' names.",
        "If a contribution from an edited volume is being cited, then the authors' names are given first.",
        "The abbreviation pp. should be used for page numbers in books.",
    ])

    with patch.object(
        service._strategy, "translate", new_callable=AsyncMock
    ) as mock_translate:
        mock_translate.return_value = [
            "縮寫 (ed.) 或 (eds.) 應跟在編輯姓名後面。",
            "引用編輯書卷中的文章時，先列出作者姓名。",
            "書籍頁碼應使用縮寫 pp.。",
        ]
        result = await service.translate_document(docx_content, "test.docx")

    assert len(result.paragraphs) == 3
    assert "ed." in result.paragraphs[0].original
    assert "ed." in result.paragraphs[0].translated or "eds." in result.paragraphs[0].translated
    assert "contribution" in result.paragraphs[1].original
    assert "引用" in result.paragraphs[1].translated
    assert "pp." in result.paragraphs[2].original
    assert "pp." in result.paragraphs[2].translated


@pytest.mark.asyncio
async def test_heading_separated_groups_translated_independently(service):
    """Groups split by headings are translated in parallel, each 1:1 mapped."""
    doc = Document()
    doc.add_paragraph("Introduction paragraph.")
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("Method details here.")
    buf = BytesIO()
    doc.save(buf)
    docx_with_heading = buf.getvalue()

    call_count = 0

    async def mock_translate(texts):
        nonlocal call_count
        call_count += 1
        if texts == ["Introduction paragraph."]:
            return ["介紹段落。"]
        if texts == ["Methods"]:
            return ["方法"]
        if texts == ["Method details here."]:
            return ["方法細節在此。"]
        raise ValueError(f"Unexpected texts: {texts}")

    with patch.object(service._strategy, "translate", side_effect=mock_translate):
        result = await service.translate_document(docx_with_heading, "test.docx")

    assert len(result.paragraphs) == 3
    assert result.paragraphs[0].translated == "介紹段落。"
    assert result.paragraphs[1].translated == "方法"
    assert result.paragraphs[2].translated == "方法細節在此。"
    # 3 groups (body, heading, body) = 3 translate calls
    assert call_count == 3


@pytest.mark.asyncio
async def test_figure_and_table_paragraphs_skip_translation(service):
    """FIGURE and TABLE paragraphs should get translated='' without calling strategy."""
    from src.services.document_parser import ParsedParagraph

    parsed = [
        ParsedParagraph(text="Normal text.", style=ParagraphStyle.NORMAL),
        ParsedParagraph(text="<::chart::>", style=ParagraphStyle.FIGURE),
        ParsedParagraph(text="<table><tr><td>X</td></tr></table>", style=ParagraphStyle.TABLE),
        ParsedParagraph(text="More text.", style=ParagraphStyle.NORMAL),
    ]

    with patch.object(service, "_parser") as mock_parser, \
         patch.object(
             service._strategy, "translate", new_callable=AsyncMock
         ) as mock_translate:
        mock_parser.parse.return_value = parsed
        mock_translate.side_effect = [
            ["普通文本。"],
            ["更多文本。"],
        ]
        result = await service.translate_document(b"fake", "test.pdf")

    # Strategy should only be called for the two NORMAL groups, not for FIGURE or TABLE
    assert mock_translate.call_count == 2

    assert result.paragraphs[0].translated == "普通文本。"
    assert result.paragraphs[1].style == ParagraphStyle.FIGURE
    assert result.paragraphs[1].translated == ""
    assert result.paragraphs[2].style == ParagraphStyle.TABLE
    assert result.paragraphs[2].translated == ""
    assert result.paragraphs[3].translated == "更多文本。"
