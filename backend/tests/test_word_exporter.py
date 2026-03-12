from io import BytesIO

from docx import Document

from src.models.translation import (
    ParagraphStyle,
    TranslatedParagraph,
    TranslationDirection,
    TranslationResult,
)
from src.services.word_exporter import WordExporter


def _make_result(**kwargs):
    defaults = dict(
        filename="test.docx",
        paragraphs=[
            TranslatedParagraph(original="Hello", translated="你好"),
            TranslatedParagraph(original="World", translated="世界"),
        ],
    )
    defaults.update(kwargs)
    return TranslationResult(**defaults)


def test_export_bilingual_table_has_two_columns():
    result = _make_result()
    exporter = WordExporter()
    docx_bytes = exporter.export(result)

    doc = Document(BytesIO(docx_bytes))
    table = doc.tables[0]
    header = table.rows[0].cells
    assert header[0].text == "English (Original)"
    assert header[1].text == "中文（翻譯）"
    # Data rows
    assert table.rows[1].cells[0].text == "Hello"
    assert table.rows[1].cells[1].text == "你好"
    assert table.rows[2].cells[0].text == "World"
    assert table.rows[2].cells[1].text == "世界"


def test_export_bilingual_zh_to_en_direction():
    result = _make_result(direction=TranslationDirection.ZH_TO_EN)
    exporter = WordExporter()
    docx_bytes = exporter.export(result)

    doc = Document(BytesIO(docx_bytes))
    table = doc.tables[0]
    header = table.rows[0].cells
    assert header[0].text == "中文（原文）"
    assert header[1].text == "English (Translation)"


def test_export_bilingual_headings_are_bold():
    result = _make_result(
        paragraphs=[
            TranslatedParagraph(
                original="Title", translated="標題", style=ParagraphStyle.TITLE
            ),
            TranslatedParagraph(
                original="Body", translated="內文", style=ParagraphStyle.NORMAL
            ),
        ],
    )
    exporter = WordExporter()
    docx_bytes = exporter.export(result)

    doc = Document(BytesIO(docx_bytes))
    table = doc.tables[0]
    # Row 1 = title → bold
    assert table.rows[1].cells[0].paragraphs[0].runs[0].bold is True
    # Row 2 = normal → not bold
    assert table.rows[2].cells[0].paragraphs[0].runs[0].bold is None


def test_export_bilingual_skips_figure_and_table():
    result = _make_result(
        paragraphs=[
            TranslatedParagraph(
                original="Hello", translated="你好", style=ParagraphStyle.NORMAL
            ),
            TranslatedParagraph(
                original="<::chart::>",
                translated="",
                style=ParagraphStyle.FIGURE,
            ),
            TranslatedParagraph(
                original="<table><tr><td>X</td></tr></table>",
                translated="",
                style=ParagraphStyle.TABLE,
            ),
            TranslatedParagraph(
                original="World", translated="世界", style=ParagraphStyle.NORMAL
            ),
        ],
    )
    exporter = WordExporter()
    docx_bytes = exporter.export(result)

    doc = Document(BytesIO(docx_bytes))
    table = doc.tables[0]
    # Header + 2 data rows (figure/table skipped)
    assert len(table.rows) == 3
    assert table.rows[1].cells[0].text == "Hello"
    assert table.rows[2].cells[0].text == "World"
